# कणाद: Research Assistant – Full Reset & Refined Implementation

from fastapi import FastAPI, Request, Form
from fastapi.responses import HTMLResponse, FileResponse
from fastapi.templating import Jinja2Templates
import requests
from bs4 import BeautifulSoup
from docx import Document
import os
import tempfile
import shutil
from difflib import SequenceMatcher
import atexit

app = FastAPI(title="कणाद: Research Assistant")

templates = Jinja2Templates(directory="templates")
temp_dir = tempfile.mkdtemp()

@atexit.register
def cleanup_temp():
    if os.path.exists(temp_dir):
        shutil.rmtree(temp_dir, ignore_errors=True)

def is_similar(a, b, threshold=0.8):
    return SequenceMatcher(None, a.lower(), b.lower()).ratio() > threshold

def extract_main_finding(paper):
    return paper.get("abstract", "").strip()

def get_google_scholar_papers(topic, max_results=10):
    search_url = f"https://scholar.google.com/scholar?q={topic.replace(' ', '+')}"
    headers = {"User-Agent": "Mozilla/5.0"}
    response = requests.get(search_url, headers=headers)
    if response.status_code != 200:
        return []
    soup = BeautifulSoup(response.text, 'html.parser')
    results = soup.select(".gs_ri")[:max_results]
    papers = []
    for result in results:
        title_tag = result.select_one("h3 a")
        if not title_tag:
            continue
        title = title_tag.text.strip()
        url = title_tag.get("href")
        snippet = result.select_one(".gs_rs")
        abstract = snippet.text.strip() if snippet else "No abstract available."
        papers.append({
            "title": title,
            "authors": [{"name": "Unknown"}],
            "year": "Unknown",
            "abstract": abstract,
            "citationCount": 0,
            "url": url,
            "externalIds": {},
            "source": "Google Scholar"
        })
    return papers

def get_taylor_and_francis_papers(topic):
    return []

def get_sciencedirect_papers(topic):
    return []

def get_asce_papers(topic):
    return []

def get_wiley_papers(topic):
    return []

def get_asme_papers(topic):
    return []

def format_apa(paper):
    authors = ", ".join([a['name'] for a in paper.get('authors', [])])
    title = paper.get('title')
    year = paper.get('year')
    source = paper.get('source', 'Unknown Source')
    doi = paper.get('externalIds', {}).get('DOI', '')
    url = paper.get('url')
    return f"{authors} ({year}). {title}. {source}. https://doi.org/{doi}" if doi else f"{authors} ({year}). {title}. {source}. {url}"

def generate_doc(papers, topic):
    doc = Document()
    doc.add_heading(f"Literature Review Summary: {topic}", 0)
    findings = []
    unique = []
    citations = []
    for paper in papers:
        abstract = extract_main_finding(paper)
        if not any(is_similar(abstract, f) for f in findings):
            findings.append(abstract)
            unique.append(paper)
            citations.append(format_apa(paper))

    for idx, paper in enumerate(unique, 1):
        doc.add_heading(f"{idx}. {paper.get('title')}", level=1)
        doc.add_paragraph(f"Source: {paper.get('source')}")
        doc.add_paragraph(f"Authors: {', '.join([a['name'] for a in paper.get('authors', [])])}")
        doc.add_paragraph(f"Year: {paper.get('year')}")
        doc.add_paragraph("Main Finding:")
        doc.add_paragraph(paper.get('abstract', 'No abstract available.'))
        doc.add_paragraph("APA Citation:")
        doc.add_paragraph(format_apa(paper))
        doc.add_paragraph("Link: " + paper.get('url'))
        doc.add_page_break()

    doc.add_heading("References", level=1)
    for citation in citations:
        doc.add_paragraph(citation)

    sanitized_topic = "_".join(topic.split())
    filename = f"Literature_Review_Report_on_{sanitized_topic}.docx"
    output_path = os.path.join(temp_dir, filename)
    doc.save(output_path)
    return output_path

@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    try:
        return templates.TemplateResponse("index.html", {"request": request})
    except:
        return HTMLResponse("""
        <html><body>
        <h2>कणाद: Research Assistant</h2>
        <form action='/search' method='post'>
            <label>Enter Research Topic:</label>
            <input name='topic' type='text' required />
            <button type='submit'>Generate</button>
        </form></body></html>
        """)

@app.post("/search")
async def search_papers(request: Request, topic: str = Form(...)):
    sources = [
        get_google_scholar_papers,
        get_taylor_and_francis_papers,
        get_sciencedirect_papers,
        get_asce_papers,
        get_wiley_papers,
        get_asme_papers
    ]
    papers = []
    for source_func in sources:
        papers.extend(source_func(topic))
    output_doc = generate_doc(papers, topic)
    sanitized_topic = "_".join(topic.split())
    download_name = f"Literature_Review_Report_on_{sanitized_topic}.docx"
    return FileResponse(output_doc, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename=download_name)
